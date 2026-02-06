#!/usr/bin/env python3
"""
Resume file parsing utilities.

Goal: reliably extract text from PDF/DOCX/TXT uploads with graceful fallbacks.
"""

from __future__ import annotations

import io
import logging
from dataclasses import dataclass, field
from typing import Optional

logger = logging.getLogger(__name__)


@dataclass
class ResumeTextExtraction:
    text: str
    warnings: list[str] = field(default_factory=list)


def extract_text_from_upload(filename: str, content: bytes) -> ResumeTextExtraction:
    """
    Extract text content from an uploaded resume file.

    Supported extensions:
      - .txt
      - .pdf
      - .docx
    """
    name = (filename or "").lower()

    if name.endswith(".txt"):
        return ResumeTextExtraction(text=_decode_text(content))

    if name.endswith(".pdf"):
        return _extract_pdf_text(content)

    if name.endswith(".docx"):
        return _extract_docx_text(content)

    # Should not happen if caller validates extensions.
    return ResumeTextExtraction(text=_decode_text(content), warnings=["Unsupported extension; used raw decode fallback."])


def _decode_text(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        return content.decode("utf-8", errors="ignore")


def _extract_pdf_text(content: bytes) -> ResumeTextExtraction:
    warnings: list[str] = []
    try:
        from PyPDF2 import PdfReader
    except Exception as e:
        logger.warning(f"PyPDF2 not available: {e}")
        return ResumeTextExtraction(
            text=_decode_text(content),
            warnings=["PyPDF2 not available; used raw decode fallback."],
        )

    try:
        reader = PdfReader(io.BytesIO(content))
        parts: list[str] = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                parts.append(page_text)
            except Exception as e:
                warnings.append(f"Failed to extract text from PDF page {i + 1}: {e}")
        text = "\n\n".join(p for p in parts if p).strip()
        if not text:
            warnings.append("PDF text extraction returned empty text; file may be scanned or malformed.")
            # Best-effort fallback to decoded bytes (often useless, but avoids hard failure).
            text = _decode_text(content)
        return ResumeTextExtraction(text=text, warnings=warnings)
    except Exception as e:
        logger.info(f"PDF extraction failed; falling back to raw decode: {e}")
        return ResumeTextExtraction(
            text=_decode_text(content),
            warnings=["Failed to parse PDF; used raw decode fallback."],
        )


def _extract_docx_text(content: bytes) -> ResumeTextExtraction:
    warnings: list[str] = []
    try:
        from docx import Document
    except Exception as e:
        logger.warning(f"python-docx not available: {e}")
        return ResumeTextExtraction(
            text=_decode_text(content),
            warnings=["python-docx not available; used raw decode fallback."],
        )

    try:
        doc = Document(io.BytesIO(content))
        parts: list[str] = []
        for para in doc.paragraphs:
            txt = (para.text or "").strip()
            if txt:
                parts.append(txt)
        # Tables can contain important data (skills, dates, etc.)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = (cell.text or "").strip()
                    if cell_text:
                        parts.append(cell_text)

        text = "\n".join(parts).strip()
        if not text:
            warnings.append("DOCX extraction returned empty text; file may be image-only or malformed.")
        return ResumeTextExtraction(text=text or _decode_text(content), warnings=warnings)
    except Exception as e:
        logger.info(f"DOCX extraction failed; falling back to raw decode: {e}")
        return ResumeTextExtraction(
            text=_decode_text(content),
            warnings=["Failed to parse DOCX; used raw decode fallback."],
        )

